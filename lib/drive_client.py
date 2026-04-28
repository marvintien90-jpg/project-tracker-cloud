"""Google Drive 讀取（用 Service Account，不需要 OAuth 互動）。"""
from __future__ import annotations

import io
import re
import warnings
from typing import Any

from docx import Document
from google.oauth2.service_account import Credentials
from googleapiclient.discovery import build
from googleapiclient.http import MediaIoBaseDownload

from .config import (
    DOC_MIME, DOCX_MIME, DRIVE_SCOPES, FOLDER_MIME, GDOC_MIME,
    get_service_account_info,
)

PDF_MIME = 'application/pdf'


def get_drive_service():
    info = get_service_account_info()
    creds = Credentials.from_service_account_info(info, scopes=DRIVE_SCOPES)
    return build('drive', 'v3', credentials=creds, cache_discovery=False)


PARSEABLE_MIMES = {DOCX_MIME, DOC_MIME, GDOC_MIME, PDF_MIME}


def list_doc_files(service, folder_id: str) -> list[dict]:
    """遞迴列出資料夾底下所有 .doc / .docx / Google Doc / PDF 檔案。"""
    out: list[dict] = []
    page_token = None
    while True:
        resp = service.files().list(
            q=f"'{folder_id}' in parents and trashed=false",
            fields='nextPageToken, files(id, name, mimeType, createdTime)',
            orderBy='createdTime desc',
            pageToken=page_token,
            pageSize=200,
        ).execute()
        for item in resp.get('files', []):
            mime = item['mimeType']
            if mime in PARSEABLE_MIMES:
                out.append(item)
            elif mime == FOLDER_MIME:
                out.extend(list_doc_files(service, item['id']))
        page_token = resp.get('nextPageToken')
        if not page_token:
            break
    return out


def parse_drive_url(url: str) -> tuple[str, str] | None:
    """從 Google Drive/Docs URL 解析出 (file_id, mime_type_hint)。

    回傳 None 表示不是可識別的 Google 檔案 URL。
    """
    # Google Docs
    m = re.search(r'docs\.google\.com/document/d/([A-Za-z0-9_-]+)', url)
    if m:
        return m.group(1), GDOC_MIME

    # Google Drive file (直接連結)
    m = re.search(r'drive\.google\.com/file/d/([A-Za-z0-9_-]+)', url)
    if m:
        return m.group(1), ''  # 需要從 API 查 mimeType

    # drive.google.com/open?id=... 或 ?usp=sharing
    m = re.search(r'[?&]id=([A-Za-z0-9_-]+)', url)
    if m and 'drive.google.com' in url:
        return m.group(1), ''

    return None


def _download_bytes(service, file_id: str) -> io.BytesIO:
    request = service.files().get_media(fileId=file_id)
    fh = io.BytesIO()
    downloader = MediaIoBaseDownload(fh, request)
    done = False
    while not done:
        _, done = downloader.next_chunk()
    fh.seek(0)
    return fh


def _export_text(service, file_id: str) -> str:
    request = service.files().export_media(fileId=file_id, mimeType='text/plain')
    fh = io.BytesIO()
    downloader = MediaIoBaseDownload(fh, request)
    done = False
    while not done:
        _, done = downloader.next_chunk()
    fh.seek(0)
    return fh.read().decode('utf-8', errors='ignore')


def extract_text(service, file_id: str, mime_type: str, filename: str) -> str:
    """根據檔案類型抽出純文字內容。

    .docx        → 用 python-docx 讀
    Google Doc   → export 成純文字
    .doc (舊版)  → 先複製成 Google Doc 再 export，最後刪除暫存
    """
    if mime_type == DOCX_MIME:
        fh = _download_bytes(service, file_id)
        doc = Document(fh)
        return '\n'.join(p.text for p in doc.paragraphs if p.text.strip())

    if mime_type == GDOC_MIME:
        return _export_text(service, file_id)

    if mime_type == DOC_MIME:
        fh = _download_bytes(service, file_id)
        raw = fh.read()
        fh.seek(0)

        # 方法一：python-docx（許多 .doc 其實是 .docx 改名）
        try:
            doc = Document(io.BytesIO(raw))
            text = '\n'.join(p.text for p in doc.paragraphs if p.text.strip())
            if text.strip():
                return text
        except Exception:
            pass

        # 方法二：olefile 直讀舊版 Word Binary Format
        try:
            import olefile
            ole = olefile.OleFileIO(io.BytesIO(raw))
            # Word 文字存在 Table stream（UTF-16 LE）
            for stream_name in ('WordDocument', '1Table', '0Table'):
                if ole.exists(stream_name):
                    data = ole.openstream(stream_name).read()
                    # 嘗試 UTF-16 LE 解碼後清理非文字字元
                    try:
                        decoded = data.decode('utf-16-le', errors='ignore')
                        cleaned = re.sub(r'[\x00-\x08\x0b\x0c\x0e-\x1f]', ' ', decoded)
                        cleaned = re.sub(r' {4,}', '\n', cleaned).strip()
                        if len(cleaned) > 50:  # 有實質內容才採用
                            return cleaned
                    except Exception:
                        continue
        except ImportError:
            pass
        except Exception:
            pass

        # 方法三：複製成 Google Doc 再 export（需要 Service Account 有儲存空間）
        try:
            copied = service.files().copy(
                fileId=file_id,
                body={'mimeType': GDOC_MIME, 'name': f'__tmp__{filename}'},
            ).execute()
        except Exception as copy_err:
            err_str = str(copy_err)
            if 'storageQuotaExceeded' in err_str:
                raise RuntimeError(
                    f'無法解析「{filename}」：三種讀取方式均失敗。'
                    '建議：① 點下方「清理 Service Account 暫存空間」釋放空間後重試，'
                    '或 ② 將檔案在 Google Drive 中另存為 Google 文件格式。'
                )
            raise
        gdoc_id = copied['id']
        try:
            return _export_text(service, gdoc_id)
        finally:
            try:
                service.files().delete(fileId=gdoc_id).execute()
            except Exception:
                pass

    if mime_type == PDF_MIME:
        try:
            import pdfplumber
        except ImportError:
            raise RuntimeError('缺少 pdfplumber 套件，請執行 pip install pdfplumber')
        fh = _download_bytes(service, file_id)
        text_parts: list[str] = []
        with warnings.catch_warnings():
            warnings.simplefilter('ignore')
            with pdfplumber.open(fh) as pdf:
                for page in pdf.pages:
                    t = page.extract_text()
                    if t:
                        text_parts.append(t)
        return '\n'.join(text_parts)

    raise ValueError(f'不支援的檔案類型：{mime_type}')


def extract_text_from_url(service, url: str) -> tuple[str, str]:
    """從任意 URL 取得文字內容。

    Google Drive/Docs URL → 用 Drive API 讀取。
    其他網址 → 用 requests + BeautifulSoup 抓取。

    回傳 (text, source_name)。
    """
    parsed = parse_drive_url(url)
    if parsed:
        file_id, mime_hint = parsed
        if not mime_hint:
            meta = service.files().get(fileId=file_id, fields='name,mimeType').execute()
            mime_hint = meta['mimeType']
            filename = meta['name']
        else:
            try:
                meta = service.files().get(fileId=file_id, fields='name').execute()
                filename = meta['name']
            except Exception:
                filename = file_id
        text = extract_text(service, file_id, mime_hint, filename)
        return text, filename

    # 一般網址
    try:
        import requests
        from bs4 import BeautifulSoup
    except ImportError:
        raise RuntimeError('缺少 requests / beautifulsoup4，請執行 pip install requests beautifulsoup4')

    headers = {'User-Agent': 'Mozilla/5.0 (compatible; ProjectTracker/1.0)'}
    resp = requests.get(url, headers=headers, timeout=20)
    resp.raise_for_status()
    soup = BeautifulSoup(resp.text, 'html.parser')
    for tag in soup(['script', 'style', 'nav', 'footer', 'header']):
        tag.decompose()
    text = soup.get_text(separator='\n', strip=True)
    source_name = soup.title.string.strip() if soup.title else url
    return text, source_name


# ============================================================
# Service Account Drive 空間管理
# ============================================================

def get_service_account_quota(service) -> dict[str, int]:
    """回傳 Service Account 自身的 Drive 儲存空間資訊（bytes）。"""
    about = service.about().get(fields='storageQuota').execute()
    quota = about.get('storageQuota', {})
    return {
        'used': int(quota.get('usage', 0)),
        'limit': int(quota.get('limit', 0)),
    }


def list_service_account_files(service) -> list[dict[str, Any]]:
    """列出 Service Account 自己 Drive 中的所有檔案（含大小）。"""
    out: list[dict[str, Any]] = []
    page_token = None
    while True:
        resp = service.files().list(
            q="'me' in owners and trashed=false",
            fields='nextPageToken, files(id, name, mimeType, size, createdTime)',
            spaces='drive',
            pageSize=200,
            pageToken=page_token,
        ).execute()
        out.extend(resp.get('files', []))
        page_token = resp.get('nextPageToken')
        if not page_token:
            break
    return out


def cleanup_service_account_drive(service, delete_all: bool = False) -> dict[str, Any]:
    """清理 Service Account Drive 中的暫存/垃圾檔案。

    delete_all=False：只刪除 __tmp__ 開頭的暫存檔
    delete_all=True ：刪除所有 Service Account 自己建立的檔案
    （Google Sheets / 使用者資料完全不受影響）

    回傳：{'deleted': N, 'freed_bytes': N, 'errors': [...]}
    """
    files = list_service_account_files(service)
    deleted = 0
    freed = 0
    errors: list[str] = []

    for f in files:
        name = f.get('name', '')
        should_delete = delete_all or name.startswith('__tmp__')
        if not should_delete:
            continue
        try:
            service.files().delete(fileId=f['id']).execute()
            deleted += 1
            freed += int(f.get('size', 0))
        except Exception as e:
            errors.append(f'{name}: {e}')

    # 同時清空垃圾桶
    try:
        service.files().emptyTrash().execute()
    except Exception:
        pass

    return {'deleted': deleted, 'freed_bytes': freed, 'errors': errors}
